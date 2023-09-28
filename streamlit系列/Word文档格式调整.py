from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
#设置对象居中、对齐等
from docx.enum.text import WD_TAB_ALIGNMENT,WD_TAB_LEADER
#设置制表符等
from docx.shared import Inches
#设置图像大小
from docx.shared import Pt, RGBColor
# 字号，颜色
from docx.oxml.ns import qn
# 中文字体
from docx.shared import Length
#设置宽度
import re #正则表达式
from docx.enum.style import WD_STYLE_TYPE #读取各部分名称
import tempfile
import os
import streamlit as st

# 创建一个Streamlit应用
st.title("更改Word文档格式！")

# 创建一个Streamlit文件上传小部件，允许用户上传Word文件
uploaded_file = st.file_uploader("上传Word文件", type=["docx"])

st.write("请输入各部分的字体字号")
ziti1 = st.text_input("一级标题字体", value='黑体', key="ziti1")
zihao1 = st.text_input("一级标题字号", value=18, key="zihao1")
zihao1 = int(zihao1)

ziti2 = st.text_input("二级标题字体", value='方正仿宋_GB2312', key="ziti2")
zihao2 = st.text_input("二级标题字号", value=16, key="zihao2")
zihao2 = int(zihao2)

ziti3 = st.text_input("三级标题字体", value='宋体', key="ziti3")
zihao3 = st.text_input("三级标题字号", value=14, key="zihao3")
zihao3 = int(zihao3)

ziti4 = st.text_input("四级标题字体", value='楷体', key="ziti4")
zihao4 = st.text_input("四级标题字号", value=12, key="zihao4")
zihao4 = int(zihao4)

ziti5 = st.text_input("正文字体", value='仿宋', key="ziti5")
zihao5 = st.text_input("正文字号", value=12, key="zihao5")
zihao5 = int(zihao5)

ziti6 = st.text_input("题注字体", value='微软雅黑', key="ziti6")
zihao6 = st.text_input("题注字号", value=8, key="zihao6")
zihao6 = int(zihao6)

st.write('一级标题:', ziti1, zihao1, '二级标题:', ziti2, zihao2, '三级标题:', ziti3, zihao3)
st.write( '四级标题:', ziti4,zihao4, '正文样式:', ziti5, zihao5, '题注样式:', ziti6, zihao6)

# 如果没有上传文件，停止应用
if uploaded_file is None:
    st.stop()

if st.button("导出Word文件"):
    # 创建一个临时文件以保存导出的Word文档
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        temp_file_path = temp_file.name

        # 打开上传的Word文档
        uploaded_doc = Document(uploaded_file)

        # 在导出Word文档中进行一些操作（这里示例中没有具体操作，仅保存上传的文档）
        # 这里你可以根据需要添加你的操作
        for 段落 in uploaded_doc.paragraphs:
            if 段落.style.name == 'Heading 1':
                for 块 in 段落.runs:
                    块.font.bold = True  # 加粗
                    块.font.italic = False  # 斜体
                    块.font.underline = False  # 下划线
                    块.font.strike = False  # 删除线
                    块.font.shadow = False  # 阴影
                    块.font.size = Pt(zihao1)
                    块.font.color.rgb = RGBColor(0, 0, 0)
                    # 颜色
                    块.font.name = 'Arial'
                    # 英文字体设置
                    块._element.rPr.rFonts.set(qn('w:eastAsia'), ziti1)
                    # 设置中文字体
        for 段落 in uploaded_doc.paragraphs:
            if 段落.style.name == 'Heading 2':
                for 块 in 段落.runs:
                    块.font.bold = True  # 加粗
                    块.font.italic = False  # 斜体
                    块.font.underline = False  # 下划线
                    块.font.strike = False  # 删除线
                    块.font.shadow = False  # 阴影
                    块.font.size = Pt(zihao2)
                    块.font.color.rgb = RGBColor(0, 0, 0)
                    # 颜色
                    块.font.name = 'Arial'
                    # 英文字体设置
                    块._element.rPr.rFonts.set(qn('w:eastAsia'), ziti2)
                    # 设置中文字体
        for 段落 in uploaded_doc.paragraphs:
            if 段落.style.name == 'Heading 3':
                for 块 in 段落.runs:
                    块.font.bold = True  # 加粗
                    块.font.italic = False  # 斜体
                    块.font.underline = False  # 下划线
                    块.font.strike = False  # 删除线
                    块.font.shadow = False  # 阴影
                    块.font.size = Pt(zihao3)
                    块.font.color.rgb = RGBColor(0, 0, 0)
                    # 颜色
                    块.font.name = 'Arial'
                    # 英文字体设置
                    块._element.rPr.rFonts.set(qn('w:eastAsia'), ziti3)
                    # 设置中文字体
        for 段落 in uploaded_doc.paragraphs:
            if 段落.style.name == 'Heading 4':
                for 块 in 段落.runs:
                    块.font.bold = True  # 加粗
                    块.font.italic = False  # 斜体
                    块.font.underline = False  # 下划线
                    块.font.strike = False  # 删除线
                    块.font.shadow = False  # 阴影
                    块.font.size = Pt(zihao4)
                    块.font.color.rgb = RGBColor(0, 0, 0)
                    # 颜色
                    块.font.name = 'Arial'
                    # 英文字体设置
                    块._element.rPr.rFonts.set(qn('w:eastAsia'), ziti4)
                    # 设置中文字体
        for 段落 in uploaded_doc.paragraphs:
            if 段落.style.name == 'Normal':
                for 块 in 段落.runs:
                    块.font.bold = False  # 加粗
                    块.font.italic = False  # 斜体
                    块.font.underline = False  # 下划线
                    块.font.strike = False  # 删除线
                    块.font.shadow = False  # 阴影
                    块.font.size = Pt(zihao5)
                    块.font.color.rgb = RGBColor(0, 0, 0)
                    # 颜色
                    块.font.name = 'Arial'
                    # 英文字体设置
                    块._element.rPr.rFonts.set(qn('w:eastAsia'), ziti5)
                    # 设置中文字体
        for 段落 in uploaded_doc.paragraphs:
            if 段落.style.name == 'TOC Heading':
                for 块 in 段落.runs:
                    块.font.bold = True  # 加粗
                    块.font.italic = True  # 斜体
                    块.font.underline = False  # 下划线
                    块.font.strike = False  # 删除线
                    块.font.shadow = False  # 阴影
                    块.font.size = Pt(16)
                    块.font.color.rgb = RGBColor(0, 0, 0)
                    # 颜色
                    块.font.name = 'Arial'
                    # 英文字体设置
                    块._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    # 设置中文字体
        for 段落 in uploaded_doc.paragraphs:
            if 段落.style.name == 'Caption':
                for 块 in 段落.runs:
                    块.font.bold = True  # 加粗
                    块.font.italic = True  # 斜体
                    块.font.underline = False  # 下划线
                    块.font.strike = False  # 删除线
                    块.font.shadow = False  # 阴影
                    块.font.size = Pt(zihao6)
                    块.font.color.rgb = RGBColor(0, 0, 255)
                    # 颜色
                    块.font.name = 'Arial'
                    # 英文字体设置
                    块._element.rPr.rFonts.set(qn('w:eastAsia'), ziti6)

        # 保存导出的Word文档
        uploaded_doc.save(temp_file_path)

    # 提供下载按钮以下载导出的Word文件

    st.download_button("点击此处下载导出的Word文件", data=open(temp_file_path, "rb").read(), file_name="更改后的Word文档.docx")




